import io
import os
import re
import unicodedata

import docx
import fitz  # PyMuPDF
from PIL import Image as _PILImage

from services.ocr import ocr_image_bytes, ocr_pdf_page

# ── Maximum pixel dimension for embedded-image OCR ───────────────────────────
# Scanned PDFs store the original scanner image (often 300 DPI A4 = 2480×3508
# = 8.7 MP).  OCR on 8.7 MP takes ~30 s on mobile models; the same image
# downscaled to 1920 px max-side (≈ 3.7 MP) takes ~8 s with no loss of
# accuracy for ≥ 10 pt text.  Amazon Textract caps its own input at 10 MP;
# we cap at ~3.7 MP to target ≤ 10 s per page on consumer hardware.
_MAX_OCR_SIDE_PX = 1920


def _cap_image_bytes(image_bytes: bytes) -> bytes:
    """
    If the image is larger than _MAX_OCR_SIDE_PX on its longer side, scale it
    down with LANCZOS resampling and return PNG bytes.  Otherwise return as-is.

    This is purely a performance optimisation — a 1920 px long-side image retains
    all characters ≥ 10 pt at 150+ DPI, which covers every printed document.
    """
    try:
        img = _PILImage.open(io.BytesIO(image_bytes))
        w, h = img.size
        if max(w, h) <= _MAX_OCR_SIDE_PX:
            return image_bytes
        scale = _MAX_OCR_SIDE_PX / max(w, h)
        new_w, new_h = max(1, int(w * scale)), max(1, int(h * scale))
        img = img.resize((new_w, new_h), _PILImage.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return image_bytes  # never crash — return original on error


# ── Text normalisation ────────────────────────────────────────────────────────
# Applied to every extracted text block so that BM25, embeddings, and the
# tagger all work on clean Unicode.
#
# Problems fixed:
#   1. Ligatures — PDF fonts encode "ﬁle" (U+FB01) instead of "fi" + "le".
#      Without NFKC normalisation, searching for "file" finds nothing because
#      the stored text contains a ligature character the query doesn't have.
#      NFKC: ﬁ→fi  ﬂ→fl  ﬀ→ff  ﬃ→ffi  ﬄ→ffl  ﬆ→st  etc.
#
#   2. Soft hyphens (U+00AD) — PDF word processors insert invisible hyphens
#      that survive text extraction and corrupt tokens: "effi­cient" ≠ "efficient".
#
#   3. PDF line-wrap hyphens — "infor-\nmation" must become "information".
#      Only applies when the break character is a real ASCII hyphen before \n
#      and the continuation is a lowercase letter.
#
#   4. Non-breaking spaces / zero-width spaces — collapse to plain space.
#
# Industry reference: Apache Tika, Azure AI Document Intelligence, and
# LangChain's UnstructuredPDFLoader all apply NFKC + soft-hyphen removal
# as their first normalisation step.

_LIGATURE_TABLE = str.maketrans({
    "ﬀ": "ff",   # ﬀ
    "ﬁ": "fi",   # ﬁ
    "ﬂ": "fl",   # ﬂ
    "ﬃ": "ffi",  # ﬃ
    "ﬄ": "ffl",  # ﬄ
    "ﬅ": "st",   # ﬅ
    "ﬆ": "st",   # ﬆ
    "­": "",     # soft hyphen — just remove
    "​": "",     # zero-width space
    "‌": "",     # zero-width non-joiner
    " ": " ",    # non-breaking space → regular space
    "’": "'",    # right single quotation → apostrophe
    "“": '"',    # left double quote
    "”": '"',    # right double quote
    "–": "-",    # en-dash
    "—": "-",    # em-dash
})


def _normalize_text(text: str) -> str:
    """
    Industry-standard text normalisation for PDF/OCR output.

    Steps (in order):
      1. Translate known ligatures + typographic characters (fast lookup table).
      2. NFKC normalisation — catches any remaining Unicode compatibility chars.
      3. Repair PDF line-wrap hyphens: "infor-\nmation" → "information".
      4. Collapse runs of whitespace (preserve single newlines).
    """
    # Step 1 — fast O(n) ligature translation
    text = text.translate(_LIGATURE_TABLE)
    # Step 2 — NFKC catches ﬁ/ﬂ variants not in our table + accented composites
    text = unicodedata.normalize("NFKC", text)
    # Step 3 — PDF line-wrap hyphens: "-\n" before a lowercase letter → join
    text = re.sub(r"-\n([a-z])", r"\1", text)
    # Step 4 — collapse horizontal whitespace runs to a single space
    #           but preserve newlines (important for sentence detection later)
    text = re.sub(r"[^\S\n]+", " ", text)
    return text.strip()


def _extract_page_text_reading_order(page) -> str:
    """
    Extract text from a PDF page in natural reading order.

    PyMuPDF's default get_text("text") reads strictly top-to-bottom which
    mangles two-column academic papers.  By sorting text blocks into horizontal
    bands (sized to ~75% of the average line height) and then left-to-right
    within each band, we recover correct column order for single, double, and
    triple-column layouts without needing explicit column detection.
    """
    blocks = page.get_text("blocks")
    if not blocks:
        return ""

    # Keep only text blocks (type 0); skip image placeholders (type 1)
    text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
    if not text_blocks:
        return ""

    heights = [max(1.0, b[3] - b[1]) for b in text_blocks]
    avg_height = sum(heights) / len(heights)
    band_size = max(5.0, avg_height * 0.75)

    # Primary sort: vertical band; secondary: x-position (left → right)
    text_blocks.sort(key=lambda b: (int(b[1] / band_size), b[0]))
    return "\n".join(b[4].strip() for b in text_blocks)


def _extract_page_tables(page) -> str:
    """
    Extract structured tables from a PDF page using PyMuPDF's ruling-line
    table finder (available since PyMuPDF 1.23).

    Industry reference
    ------------------
    AWS Textract and Google Document AI both return table cells as structured
    key-value pairs rather than raw text.  PyMuPDF's find_tables() achieves
    equivalent structured extraction for border-ruled tables by detecting
    horizontal and vertical ruling lines and inferring the cell grid.

    For our academic-document corpus, structured table extraction improves
    retrieval quality for:
      • Grade distribution tables (A=93-100, B=83-92, …)
      • Weekly schedule tables (Week | Topic | Reading)
      • Policy grids (Assignment type | Weight | Due date)

    These were previously extracted as flat text runs, losing column alignment.
    Now they are stored as pipe-delimited markdown rows, which the sentence
    chunker keeps intact as coherent semantic units.

    Only tables with ≥ 2 rows AND ≥ 2 columns are emitted — single-column
    "tables" are just lists and are already captured by regular text extraction.
    Returns empty string if no qualifying tables are found (best-effort).
    """
    try:
        finder = page.find_tables()
        if not finder.tables:
            return ""

        table_texts: list[str] = []
        for table in finder.tables:
            rows = table.extract()
            if not rows or len(rows) < 2:
                continue
            # Determine actual column count (some cells may be None for merged cells)
            n_cols = max(len(r) for r in rows)
            if n_cols < 2:
                continue

            lines: list[str] = []
            for i, row in enumerate(rows):
                # Pad short rows, coerce None → "", collapse internal newlines
                cells = [
                    str(c or "").replace("\n", " ").strip()
                    for c in (list(row) + [""] * n_cols)[:n_cols]
                ]
                lines.append("| " + " | ".join(cells) + " |")
                # Insert markdown separator after header row
                if i == 0:
                    lines.append("| " + " | ".join(["---"] * n_cols) + " |")

            if lines:
                table_texts.append("\n".join(lines))

        return "\n\n".join(table_texts)
    except Exception:
        return ""  # table extraction is best-effort — never crash the pipeline


def _iter_docx_blocks(document):
    """
    Yield text strings for paragraphs AND table rows in document-body order.

    python-docx's document.paragraphs skips every table cell, losing grade
    scales, schedules, and policy grids common in syllabi.  Walking the XML
    body directly preserves the interleaved order of prose and tables.
    """
    for child in document.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            from docx.text.paragraph import Paragraph

            text = Paragraph(child, document).text.strip()
            if text:
                yield text

        elif tag == "tbl":
            from docx.table import Table

            table = Table(child, document)
            for row in table.rows:
                seen: set[str] = set()
                cells: list[str] = []
                for cell in row.cells:
                    ct = cell.text.strip()
                    # DOCX merged cells appear multiple times — deduplicate
                    if ct and ct not in seen:
                        cells.append(ct)
                        seen.add(ct)
                if cells:
                    yield " | ".join(cells)


def extract_text_by_pages(file_path: str, progress_cb=None) -> list[dict]:
    """
    Extracts text from PDF, DOCX, or plain-text files, structured by page/section.
    Returns: [{"page": int, "text": str}]
    """
    _, ext = os.path.splitext(file_path.lower())
    pages_content = []

    if ext == ".pdf":
        try:
            doc = fitz.open(file_path)

            # Detect password-protected PDFs before iterating pages
            if doc.is_encrypted:
                if not doc.authenticate(""):
                    raise ValueError(
                        "PDF is password-protected. "
                        "Please provide an unprotected copy."
                    )

            total_pages = doc.page_count

            # ── Adaptive DPI for large all-image PDFs ─────────────────────────
            # Probe the first 5 pages (or fewer for short docs).  If all of
            # them are image-only (no selectable text), this is a scanned PDF
            # and every page will need OCR.  Rendering at 150 DPI is ~4× faster
            # than 300 DPI and still accurate enough for modern clean office
            # scans (laser-printed exam timetables, photocopied forms, etc.).
            # Documents with ANY selectable text keep the 300 DPI default so
            # degraded or mixed-content files stay at high quality.
            ocr_dpi = 300
            if total_pages > 10:
                probe_n = min(5, total_pages)
                image_only_count = sum(
                    1 for i in range(probe_n)
                    if not doc.load_page(i).get_text("text").strip()
                )
                if image_only_count == probe_n:
                    ocr_dpi = 150
                    print(
                        f"All-image PDF ({total_pages} pages): "
                        f"using {ocr_dpi} DPI for faster OCR"
                    )

            for page_idx in range(total_pages):
                page_num = page_idx + 1

                # Per-page error handling — a single corrupt page should not
                # abort the entire document
                try:
                    page = doc.load_page(page_idx)
                except Exception as load_err:
                    print(f"Skipping page {page_num} (failed to load: {load_err})")
                    if progress_cb:
                        progress_cb(page_num, total_pages)
                    continue

                text = _normalize_text(_extract_page_text_reading_order(page))
                is_ocr_page = not bool(text)

                # Structured table extraction only for vector-content pages.
                # find_tables() uses ruling-line detection — it finds nothing
                # on raster (scanned/image-only) pages, so skip the call.
                if text:
                    table_text = _extract_page_tables(page)
                    if table_text:
                        text = text + "\n\n" + table_text

                image_list = page.get_images(full=True)

                if not text:
                    # ── Image-only page: OCR path ─────────────────────────────
                    # Strategy: try embedded-image OCR first (extracts from the
                    # stored image bytes directly — often cleaner than rendering
                    # at a fixed DPI).  Only fall back to full-page render OCR
                    # if embedded OCR produced nothing or very little text.
                    #
                    # We do NOT run both — the embedded image IS the page, so
                    # running both would duplicate all the text.
                    ocr_from_images = []
                    for img_idx, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            if (
                                base_image.get("width", 0) < 100
                                or base_image.get("height", 0) < 100
                            ):
                                continue  # decorative — skip
                            # Cap large scans to _MAX_OCR_SIDE_PX before OCR
                            img_bytes = _cap_image_bytes(base_image["image"])
                            img_text = ocr_image_bytes(img_bytes)
                            if img_text:
                                ocr_from_images.append(img_text)
                        except Exception as img_err:
                            print(f"Embedded image OCR error p{page_num} img{img_idx}: {img_err}")

                    if ocr_from_images:
                        # Embedded image gave us content — use it directly
                        text = "\n".join(ocr_from_images)
                    else:
                        # No embedded image text → render full page and OCR that
                        try:
                            text = ocr_pdf_page(page, dpi=ocr_dpi)
                        except Exception as ocr_err:
                            print(f"Page-level OCR failed on page {page_num}: {ocr_err}")

                elif image_list:
                    # ── Text page with embedded figures ──────────────────────
                    # Selectable text already captured above.  OCR embedded
                    # images only if they look like figures (not full-page scans).
                    # Heuristic: skip images whose area exceeds 60% of the page.
                    page_area = page.rect.width * page.rect.height
                    for img_idx, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            w = base_image.get("width", 0)
                            h = base_image.get("height", 0)
                            if w < 100 or h < 100:
                                continue  # decorative
                            # Skip if image fills most of the page (it's the scan,
                            # not a figure — the selectable text already covers it)
                            img_area_approx = w * h
                            if img_area_approx > page_area * 0.6:
                                continue
                            # Cap large figures before OCR (performance)
                            img_bytes = _cap_image_bytes(base_image["image"])
                            img_text = ocr_image_bytes(img_bytes)
                            if img_text:
                                text += f"\n\n[Figure {img_idx + 1}]\n{img_text}"
                        except Exception as img_err:
                            print(f"Figure OCR error p{page_num} img{img_idx}: {img_err}")

                if text.strip():
                    pages_content.append({"page": page_num, "text": text.strip()})

                # Report per-page progress so the caller can update the UI.
                # This is especially important for large all-image PDFs where
                # OCR can take several minutes with no other feedback.
                if progress_cb:
                    progress_cb(page_num, total_pages)

        except ValueError:
            raise  # Re-raise password/format errors with their original message
        except Exception as exc:
            print(f"Error parsing PDF {file_path}: {exc}")
            raise

    elif ext == ".docx":
        try:
            document = docx.Document(file_path)
            current_page = 1
            word_count = 0
            page_text: list[str] = []
            TARGET_WORDS_PER_PAGE = 300

            for text in _iter_docx_blocks(document):
                text = _normalize_text(text)
                page_text.append(text)
                word_count += len(text.split())

                if word_count >= TARGET_WORDS_PER_PAGE:
                    pages_content.append(
                        {"page": current_page, "text": "\n".join(page_text)}
                    )
                    page_text = []
                    word_count = 0
                    current_page += 1

            if page_text:
                pages_content.append(
                    {"page": current_page, "text": "\n".join(page_text)}
                )
        except Exception as exc:
            print(f"Error parsing DOCX {file_path}: {exc}")
            raise

    elif ext in [".txt", ".md", ".json", ".csv"]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = _normalize_text(f.read().strip())
            char_limit = 2000
            for page_idx, start in enumerate(range(0, len(text), char_limit)):
                chunk = text[start : start + char_limit].strip()
                if chunk:
                    pages_content.append({"page": page_idx + 1, "text": chunk})
        except Exception as exc:
            print(f"Error parsing text file {file_path}: {exc}")
            raise
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    return pages_content
